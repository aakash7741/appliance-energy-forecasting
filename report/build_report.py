#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
FIG = OUT / "figures"
TABLE = OUT / "tables"
REPORT = ROOT / "report"
BLUE = "17324D"
ACCENT = "2D6CDF"
TEAL = "1BA3A3"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
ORANGE = "E67E22"
GRAY = "667085"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=50, start=80, bottom=50, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = code
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.60)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(8.6)
    normal.font.color.rgb = RGBColor.from_string(BLUE)
    normal.paragraph_format.space_after = Pt(2.4)
    normal.paragraph_format.line_spacing = 1.04

    for name, size, before, after, color in (
        ("Heading 1", 13.5, 2, 4, BLUE),
        ("Heading 2", 10.2, 4, 2, ACCENT),
        ("Heading 3", 9.0, 3, 1, BLUE),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "APPLIANCE ENERGY FORECASTING  /  REPRODUCIBLE MODEL COMPARISON"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(6.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(GRAY)
    p_pr = header._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:color"), "D5DCE5")
    border.append(bottom); p_pr.append(border)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("21 JULY 2026   •   ")
    run.font.size = Pt(6.5); run.font.color.rgb = RGBColor.from_string(GRAY)
    add_field(footer, "PAGE")
    footer.add_run(" / 8")
    for run in footer.runs:
        run.font.name = "Aptos"; run.font.size = Pt(6.5); run.font.color.rgb = RGBColor.from_string(GRAY)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = "Aptos Display"; run.font.size = Pt(23); run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(7)
        run = s.add_run(subtitle)
        run.font.name = "Aptos"; run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string(ACCENT)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold_lead=None, size=None, color=None, after=None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = keep
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead); r1.bold = True
        r2 = p.add_run(text[len(bold_lead):])
        runs = (r1, r2)
    else:
        runs = (p.add_run(text),)
    for run in runs:
        if size: run.font.size = Pt(size)
        if color: run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_callout(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table(table, [6.9])
    cell = table.cell(0, 0); set_cell_shading(cell, fill); set_cell_margins(cell, 90, 130, 90, 130)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "  "); r.bold = True; r.font.size = Pt(7.2); r.font.color.rgb = RGBColor.from_string(ACCENT)
    r = p.add_run(text); r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(BLUE)
    return table


def add_table(doc, headers, rows, widths, font_size=7.0, header_fill=PALE, align_numbers=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]; cell.text = str(header); set_cell_shading(cell, header_fill)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    set_fixed_table(table, widths)
    for ridx, row in enumerate(table.rows):
        row.height = None
        for cidx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if align_numbers and cidx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.name = "Aptos"; run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(BLUE)
                    if ridx == 0: run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc, filename, width, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(FIG / filename), width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(6.7); r.font.color.rgb = RGBColor.from_string(GRAY)


def page_break(doc):
    doc.add_page_break()


def fmt(value, digits=2):
    return f"{float(value):.{digits}f}"


def main():
    metrics = pd.read_csv(TABLE / "model_metrics.csv")
    stationarity = pd.read_csv(TABLE / "stationarity_tests.csv")
    quality = json.loads((OUT / "data_quality.json").read_text())
    problem = json.loads((OUT / "problem_definition.json").read_text())
    summary = json.loads((OUT / "run_summary.json").read_text())
    ljung = pd.read_csv(TABLE / "sarimax_ljung_box.csv", index_col=0)
    importance = pd.read_csv(TABLE / "feature_importance.csv")
    m = metrics.set_index("Model")

    weekly = m.loc["Weekly seasonal naive"]
    sarcal = m.loc["SARIMAX calendar"]
    sarcond = m.loc["SARIMAX conditional"]
    hgb_all = m.loc["HGB all covariates (conditional)"]
    hgb_lag = m.loc["HGB lag+time"]
    chronos = m.loc["Chronos-T5 tiny zero-shot"]
    sar_gain = (weekly.RMSE - sarcal.RMSE) / weekly.RMSE * 100
    hgb_margin = (sarcal.RMSE - hgb_all.RMSE) / sarcal.RMSE * 100

    doc = Document()
    configure_document(doc)

    # PAGE 1 — Executive summary
    add_title(doc, "Forecasting appliance energy use", "Hourly smart-home demand • 24-hour horizon • benchmarks, SARIMAX, HGB and Chronos-T5")
    add_callout(doc, "Decision", f"Recommend calendar-only SARIMAX: RMSE {sarcal.RMSE:.2f} Wh, {sar_gain:.1f}% below the strongest benchmark, with 94.9% coverage for nominal 95% intervals. Conditional HGB is only {hgb_margin:.1f}% better in RMSE and requires future inputs.")
    add_heading(doc, "Executive summary", 1)
    add_para(doc, "This study forecasts hourly appliance energy from the UCI low-energy-house data. Six 10-minute Wh readings were summed to each hour; sensor and weather readings were averaged. The resulting 3,290-hour series is complete, strongly spiked, and patterned by hour and week. The task is a 24-hour-ahead forecast evaluated over the final 14 days as 14 rolling daily origins. This respects chronology and makes every model solve the same operational problem.")
    add_para(doc, "Weekly seasonal naive is the strongest of naive, daily seasonal naive, weekly seasonal naive and drift (RMSE 474.70 Wh), implying that occupancy routines repeat more reliably week-to-week than from the immediately preceding day. Yet a constant historical mean scores 436.34 Wh: much of the apparent seasonal signal consists of irregular bursts whose exact timing is hard to predict.")
    add_para(doc, "The lowest RMSE is conditional HGB (378.17 Wh), followed almost indistinguishably by calendar-only SARIMAX (380.34 Wh). HGB reduces MAE, but its all-covariate result uses realised future sensors and weather. Chronos-T5 tiny is competitive with the seasonal baselines (RMSE 460.51 Wh) but biased low by 196.38 Wh and its 90% interval covers only 66.96%, so zero-shot simplicity does not compensate for poor peak calibration here.")
    top = metrics.sort_values("RMSE").head(7)
    add_table(doc, ["Model", "RMSE", "MAE", "MASE", "Bias"], [[r.Model.replace(" (conditional)", "*").replace(" zero-shot", "")[:31], fmt(r.RMSE), fmt(r.MAE), fmt(r.MASE, 3), fmt(r.Bias)] for _, r in top.iterrows()], [2.8, 1.0, 1.0, 0.9, 1.0], 6.8)
    add_para(doc, "*Conditional: realised future covariates are supplied. Metrics are Wh except dimensionless MASE.", size=6.6, color=GRAY)
    add_callout(doc, "Practical recommendation", "Deploy calendar-only SARIMAX as the default because it is honest at forecast origin, nearly point-optimal, interpretable, inexpensive, and produces calibrated uncertainty. Monitor residual seasonality and retrain as household routines change.", fill=LIGHT)

    # PAGE 2 — Data and EDA
    page_break(doc)
    add_heading(doc, "1. Dataset and preprocessing", 1)
    add_para(doc, "The public dataset contains 19,735 observations from 11 January to 27 May 2016, sampled every 10 minutes in one Belgian low-energy house [1,2]. The target is appliance energy (`Appliances`, Wh). Covariates comprise light energy; temperature and relative humidity in nine indoor areas; outdoor-side sensors; and airport temperature, humidity, pressure, wind, visibility and dew point. The random variables `rv1` and `rv2` were excluded because they are deliberately non-predictive controls.")
    add_para(doc, "Timestamp parsing produced a unique, strictly increasing index. There are zero missing cells, missing 10-minute timestamps, duplicate timestamps or missing hourly timestamps. Appliance and light Wh readings are additive over time and were summed; temperatures, humidities and weather states were averaged. Summing is essential: averaging the six energy readings would change the physical quantity and shrink the target six-fold.")
    add_figure(doc, "01_eda_overview.png", 6.75, "Figure 1. Hourly energy, holdout detail, intraday profile and day-of-week means.")
    add_table(doc, ["Check", "Result", "Check", "Result"], [
        ["Raw rows", f"{quality['raw_rows']:,}", "Hourly rows", f"{quality['hourly_rows']:,}"],
        ["Missing cells", quality["raw_missing_cells"], "Missing hours", quality["missing_timestamps_hourly"]],
        ["Mean ± SD", f"{quality['hourly_target_mean_Wh']:.0f} ± {quality['hourly_target_std_Wh']:.0f} Wh", "Range", f"{quality['hourly_target_min_Wh']:.0f}–{quality['hourly_target_max_Wh']:.0f} Wh"],
    ], [1.25, 2.0, 1.25, 2.4], 6.9)
    add_para(doc, "Interpretation. The series combines a stable low base with intermittent appliance events up to 3,650 Wh. Mean demand peaks around 18:00 and is highest on Monday and Saturday, but mean–median separation confirms right-skewed bursts. A weekly or calendar-aware model is plausible, while peak timing limits achievable squared-error accuracy.", bold_lead="Interpretation. ")

    # PAGE 3 — Components and stationarity
    page_break(doc)
    add_heading(doc, "2. Components, seasonality and stationarity", 1)
    add_para(doc, "STL with period 24 separates a slowly moving level, a repeatable daily profile and a large irregular remainder. The ACF oscillates every 24 hours and remains positive at one week; empirical correlations are 0.581, 0.306 and 0.329 at lags 1, 24 and 168. Seasonality is therefore real but moderate: daily and weekly repetitions explain routine, not the occurrence or magnitude of every usage spike.")
    add_figure(doc, "02_components_acf.png", 6.72, "Figure 2. Robust STL components and ACF through eight days; daily oscillation and weekly persistence are visible.")
    stat_rows = []
    for _, r in stationarity.iterrows():
        stat_rows.append([r["Transformation"], fmt(r["ADF statistic"], 3), f"{r['ADF p-value']:.2g}", fmt(r["KPSS statistic"], 3), f"≥{r['KPSS p-value']:.2f}"])
    add_table(doc, ["Transformation", "ADF stat", "ADF p", "KPSS stat", "KPSS p"], stat_rows, [2.4, 1.1, 1.0, 1.2, 1.1], 6.8)
    add_para(doc, "ADF rejects a unit root and KPSS does not reject level stationarity even before differencing. This is not contradictory to seasonality: a stationary process may have a deterministic seasonal mean and short-memory autocorrelation. The AIC screen nonetheless selected d=1, suggesting differencing improves likelihood once ARMA errors are modelled. Daily differencing is diagnostically stationary but was not forced, avoiding unnecessary over-differencing; daily dependence is handled by seasonal AR and MA terms.")
    add_callout(doc, "Caution", "Statistical stationarity tests do not make the household invariant. The sample spans only 4.5 months, so annual seasonality, holidays, occupancy changes and concept drift cannot be learned or ruled out.", fill=LIGHT)

    # PAGE 4 — Forecasting design and benchmarks
    page_break(doc)
    add_heading(doc, "3. Forecasting problem and evaluation design", 1)
    add_table(doc, ["Design element", "Choice", "Why"], [
        ["Target", "Hourly appliance Wh", "Operational energy total; non-negative"],
        ["Horizon", "24 hours", "Day-ahead smart-home planning"],
        ["Train", "2,954 hours to 13 May 18:00", "All history before holdout"],
        ["Test", "336 hours / 14 rolling origins", "Multiple days, not one lucky split"],
        ["Primary metric", "RMSE", "Penalises costly missed peaks"],
        ["Secondary", "MAE, MAPE, sMAPE, MASE, bias", "Magnitude, scale-free accuracy, calibration"],
    ], [1.25, 2.35, 3.2], 7.0)
    add_para(doc, "At each origin, models issue 24 recursive or multi-step predictions. Between origins, the newly observed 24 target values enter history, while fitted parameters remain fixed. MASE uses the in-sample mean absolute 24-hour seasonal difference as its denominator. This rolling-origin design produces 336 commensurate errors and avoids random splits, which would leak future seasonal regimes into training [3].")
    add_heading(doc, "4. Benchmark forecasts", 1)
    bench_names = ["Mean", "Naive", "Daily seasonal naive", "Weekly seasonal naive", "Drift"]
    bench = metrics[metrics.Model.isin(bench_names)].set_index("Model").loc[bench_names].reset_index()
    add_table(doc, ["Benchmark", "RMSE", "MAE", "sMAPE %", "MASE", "Bias"], [[r.Model, fmt(r.RMSE), fmt(r.MAE), fmt(r.sMAPE_pct), fmt(r.MASE, 3), fmt(r.Bias)] for _, r in bench.iterrows()], [2.25, 0.9, 0.9, 1.0, 0.9, 1.0], 6.9)
    add_para(doc, "The naive and drift models are weakest because the last observed hour is often a spike or trough and drift extrapolates a negligible long-run slope from an unrepresentative endpoint. Daily seasonal naive improves substantially, but weekly seasonal naive is stronger (RMSE 474.70 versus 510.70 Wh). The mean model unexpectedly reaches 436.34 Wh because squared error rewards a conservative centre when burst timing is weakly repeatable.")
    add_callout(doc, "Question 1 — strongest benchmark", "Weekly seasonal naive is strongest among naive, daily seasonal naive, weekly seasonal naive and drift. Appliance use carries weekly household-routine information beyond the daily cycle, but the mean model’s advantage and high residual variance show that seasonality is incomplete and peaks are irregular.")
    add_heading(doc, "Fairness and leakage controls", 2)
    add_para(doc, "All lag and rolling features are shifted at least one hour. Within each 24-hour HGB path, unavailable target lags are replaced by the model’s own earlier forecasts. Hyperparameters and AIC orders use training data only. Test outcomes are used once for comparison, not for order selection. Conditional variants are labelled throughout rather than presented as deployable forecasts.")

    # PAGE 5 — SARIMAX
    page_break(doc)
    add_heading(doc, "5. SARIMAX modelling and diagnostics", 1)
    add_para(doc, "Every required non-seasonal combination p=0…6, d=0…2, q=0…6 was fit—147 likelihood models—on the latest 60 training days, with checkpoints written after each fit. Non-converged solutions were excluded before AIC selection. The winner is ARIMA(0,1,6), AIC 20,972.94. The final models add daily SARMA(1,0,1) at s=24. A calendar version uses sine/cosine hour and weekday terms plus weekend status; a conditional version adds realised future T1, RH_1, T_out, RH_out and wind speed.")
    add_table(doc, ["Model", "RMSE", "MAE", "MASE", "Bias", "95% cov."], [
        ["Weekly seasonal naive", fmt(weekly.RMSE), fmt(weekly.MAE), fmt(weekly.MASE, 3), fmt(weekly.Bias), "—"],
        ["SARIMAX calendar", fmt(sarcal.RMSE), fmt(sarcal.MAE), fmt(sarcal.MASE, 3), fmt(sarcal.Bias), f"{sarcal.Coverage_pct:.1f}%"],
        ["SARIMAX conditional", fmt(sarcond.RMSE), fmt(sarcond.MAE), fmt(sarcond.MASE, 3), fmt(sarcond.Bias), f"{sarcond.Coverage_pct:.1f}%"],
    ], [2.25, 0.9, 0.9, 0.9, 0.9, 1.0], 7.0)
    add_figure(doc, "05_sarimax_residuals.png", 6.8, "Figure 3. Conditional SARIMAX in-sample residual path, residual ACF and Q-Q plot.")
    add_para(doc, f"The calendar SARIMAX reduces RMSE by {sar_gain:.1f}% versus weekly seasonal naive and its 95% interval coverage is 94.94%, close to nominal. Residual ACF values are individually small, but Ljung–Box tests at 24 and 48 lags give p={ljung.iloc[0].lb_pvalue:.3f} and {ljung.iloc[1].lb_pvalue:.3f}: weak dependence remains. The Q-Q plot shows heavy, right-skewed tails, consistent with missed appliance bursts; Gaussian intervals can cover well overall while understating event-specific tail risk.")
    add_callout(doc, "Question 2 — does SARIMAX improve?", "Yes for calendar SARIMAX, materially. Daily seasonality and autocorrelation are mostly captured, but not completely: Ljung–Box rejects white noise and tail residuals are non-Gaussian. The conditional exogenous version worsens RMSE to 456.91 Wh, so realised sensors do not add stable linear predictive information after dynamics and calendar structure.")
    add_para(doc, "AIC limitation. Searching the latest 60 days controls computation and focuses on the contemporary regime, but it can select a different order than a full-history search. Seasonal orders were fixed parsimoniously rather than exhaustively searched; expanding P,D,Q and using time-series cross-validation are priority sensitivity checks.", bold_lead="AIC limitation. ")

    # PAGE 6 — Feature model
    page_break(doc)
    add_heading(doc, "6. Feature-based regression and covariates", 1)
    add_para(doc, "A histogram gradient boosting regressor predicts log(1+Wh), then forecasts recursively. Three nested specifications separate genuine information gains: (i) deployable lags, rolling summaries and calendar features; (ii) the same plus realised future airport/outdoor weather; and (iii) all weather, indoor temperature/humidity and lights. Lags are 1, 2, 3, 24, 25, 48 and 168 hours; shifted rolling mean and standard deviation windows are 3, 24 and 168 hours.")
    add_table(doc, ["Feature model", "Availability", "RMSE", "MAE", "Δ RMSE vs prior"], [
        ["Lag + time", "Known/recursive", fmt(hgb_lag.RMSE), fmt(hgb_lag.MAE), "—"],
        ["+ weather", "Conditional", fmt(m.loc["HGB + weather (conditional)"].RMSE), fmt(m.loc["HGB + weather (conditional)"].MAE), f"{m.loc['HGB + weather (conditional)'].RMSE-hgb_lag.RMSE:+.2f}"],
        ["+ indoor sensors/lights", "Conditional", fmt(hgb_all.RMSE), fmt(hgb_all.MAE), f"{hgb_all.RMSE-m.loc['HGB + weather (conditional)'].RMSE:+.2f}"],
    ], [2.2, 1.45, 0.95, 0.95, 1.15], 6.9)
    add_figure(doc, "06_feature_importance.png", 6.35, "Figure 4. Permutation importance on the conditional holdout; lag 1 and cyclic time dominate.")
    add_para(doc, "The deployable HGB already achieves 389.16 Wh RMSE. Supplying future weather lowers RMSE by 8.28 Wh and all sensors by only another 2.71 Wh, although conditional MAE is roughly 7.5 Wh lower than lag+time. Permutation importance is dominated by lag 1, hour cosine, lag 24 and hour sine; the first sensor (RH_6) is a distant fifth. Thus dynamics and time of day carry most useful signal, and the broad sensor panel has low marginal value.")
    add_callout(doc, "Questions 3 and 5 — feature value and availability", "Feature engineering improves strongly over the weekly benchmark, but most of the gain comes from lags and time. Calendar terms are known. Appliance lags are known only up to the origin and must be generated recursively thereafter. Future indoor temperature, humidity, lights and realised weather are not genuinely known; using their test values creates a conditional forecast, not a deployable one. Weather forecasts could be substituted, with their uncertainty propagated.")

    # PAGE 7 — Foundation model and comparison
    page_break(doc)
    add_heading(doc, "7. Foundation model and common evaluation", 1)
    add_para(doc, "Chronos-T5 tiny is a pretrained probabilistic time-series transformer that scales and quantises values, then models tokens with a T5 architecture [4]. It is used zero-shot: 512 hours of context, 40 sampled paths per origin, median point forecast and 5th–95th percentile interval. No target-specific training or future covariates are used. This makes the comparison operationally honest but gives the small univariate checkpoint no household-specific adaptation.")
    full_rows = []
    for _, r in metrics.sort_values("RMSE").iterrows():
        coverage = "—" if pd.isna(r.Coverage_pct) else f"{r.Coverage_pct:.1f}"
        name = r.Model.replace(" (conditional)", "*").replace(" zero-shot", "")
        full_rows.append([name, fmt(r.RMSE), fmt(r.MAE), fmt(r.sMAPE_pct), fmt(r.MASE, 3), fmt(r.Bias), coverage])
    add_table(doc, ["Model", "RMSE", "MAE", "sMAPE", "MASE", "Bias", "Cov.%"], full_rows, [2.6, 0.75, 0.75, 0.75, 0.7, 0.8, 0.7], 6.35)
    add_para(doc, "*Conditional forecast. Coverage corresponds to each model’s stated interval (95% SARIMAX; 90% Chronos).", size=6.3, color=GRAY)
    add_figure(doc, "07_model_comparison.png", 6.65, "Figure 5. Common RMSE and MASE comparison; all learned models beat the strongest seasonal benchmark except conditional SARIMAX.")
    add_callout(doc, "Question 4 — does the foundation model justify itself?", f"No. Chronos RMSE is {chronos.RMSE:.2f} Wh: better than weekly seasonal naive but worse than SARIMAX and HGB. Its low sMAPE reflects conservative forecasts, yet bias is {chronos.Bias:.2f} Wh and 90% coverage is only {chronos.Coverage_pct:.1f}%. The tiny checkpoint under-predicts bursts; larger/contextual foundation models or fine-tuning might help, but the observed result does not justify extra dependencies or latency.")

    # PAGE 8 — Forecasts, recommendation, limitations, references
    page_break(doc)
    add_heading(doc, "8. Final 24-hour forecasts, recommendation and limitations", 1)
    add_figure(doc, "04_final_forecasts.png", 5.35, "Figure 6. Every model at the final rolling origin versus the subsequent observed 24 hours; shaded bands are available prediction intervals.")
    add_para(doc, "After the final origin, observed demand averaged 750.83 Wh and peaked at 1,860 Wh at 11:00. Calendar SARIMAX is the best of the plotted final-day forecasts (RMSE 420.16 Wh) but still smooths the midday bursts; HGB-all peaks at 1,142 Wh and Chronos at only 660 Wh. Daily seasonal naive imports a false 2,970 Wh evening peak. These errors reinforce the aggregate diagnostics: models learn routine level and timing, while discrete occupant-driven events remain the dominant uncertainty.")
    add_callout(doc, "Question 6 — final recommendation", "Choose calendar-only SARIMAX. Its accuracy is essentially tied with conditional HGB, it is interpretable, its inputs are truly known, its uncertainty is calibrated, and CPU deployment is simple. HGB lag+time is a credible alternative when lower MAE and nonlinear effects matter; conditional HGB should be used only with honest covariate forecasts.")
    add_para(doc, "Limitations and improvements. One house, 4.5 months and no occupancy/appliance-state labels limit external validity and peak prediction. Future work should use probabilistic peak-aware losses, conformal calibration, separate weather forecasts, occupancy/calendar events, richer seasonal-order cross-validation, tuned HGB, larger or covariate-aware foundation models, and multi-home/annual data. Operational monitoring should track bias, interval coverage and drift by hour.", bold_lead="Limitations and improvements. ")
    add_heading(doc, "References", 2)
    refs = (
        "[1] Candanedo, L. (2017). Appliances Energy Prediction. UCI ML Repository. doi:10.24432/C5VC8G.   "
        "[2] Candanedo, L.M., Feldheim, V. & Deramaix, D. (2017). Data driven prediction models of energy use of appliances in a low-energy house. Energy and Buildings 140, 81–97. doi:10.1016/j.enbuild.2017.01.083.   "
        "[3] Hyndman, R.J. & Athanasopoulos, G. (2021). Forecasting: Principles and Practice, 3rd ed. OTexts.   "
        "[4] Ansari, A.F. et al. (2024). Chronos: Learning the language of time series. TMLR, arXiv:2403.07815.   "
        "[5] statsmodels (2026). SARIMAX documentation, v0.14.6."
    )
    add_para(doc, refs, size=6.3, color=GRAY, after=0)

    properties = doc.core_properties
    properties.title = "Forecasting Appliance Energy Use"
    properties.subject = "Hourly 24-hour-ahead forecasting comparison"
    properties.author = "Reproducible analysis"
    properties.keywords = "time series, appliance energy, SARIMAX, histogram gradient boosting, Chronos"

    REPORT.mkdir(parents=True, exist_ok=True)
    path = REPORT / "appliance_energy_forecasting_report.docx"
    doc.save(path)
    print(path)


if __name__ == "__main__":
    main()
